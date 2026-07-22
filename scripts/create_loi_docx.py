from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Header block ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Jason Christiansen')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('[ADDRESS LINE 1]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('[PHONE]  |  [EMAIL]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

# Date
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('[DATE]')
run.font.size = Pt(11)

# Addressee
doc.add_paragraph()
p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('VIA EMAIL')
run.bold = True
run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('Kevin [LAST NAME]')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('[COMPANY NAME]')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.space_after = Pt(0)
run = p.add_run('[ADDRESS]')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('[EMAIL]')
run.font.size = Pt(11)

# Re line
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Re: Letter of Intent for Purchase and Sale of ±62 Acres, N Mayflower Drive, Greenville, WI 54942')
run.bold = True
run.font.size = Pt(11)

# Salutation
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Dear Kevin,')
run.font.size = Pt(11)

# Body - Letter of Intent text
body = """This letter of intent, effective as of [DATE] (the "Effective Date"), constitutes an expression of interest by Jason Christiansen ("Purchaser") in purchasing certain property owned by the Estler/Barbara family ("Seller") on the general terms and conditions described herein. This LOI serves as the basis for negotiating a definitive Purchase Agreement."""

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(body)
run.font.size = Pt(11)

# --- Terms ---
def add_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

def add_term(label, value):
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(11)

add_heading('1. The Property')
p = doc.add_paragraph()
run = p.add_run('Two contiguous parcels totaling approximately 62 acres, located on N Mayflower Drive, Greenville, Outagamie County, WI 54942, further identified by Outagamie County parcel numbers:')
run.font.size = Pt(11)
add_bullet('Parcel 1: [APN] — [± acres]')
add_bullet('Parcel 2: [APN] — [± acres]')

add_heading('2. Purchase Price')
p = doc.add_paragraph()
run = p.add_run('$900,000 (Nine Hundred Thousand and 00/100 Dollars)')
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Structured as:')
run.font.size = Pt(11)

add_bullet('$40,000 Earnest Money Deposit upon execution of Purchase Agreement')
add_bullet('$860,000 Seller Carry Note paid at closing')

add_heading('3. Earnest Money Deposit')
add_bullet('$40,000 deposited within 5 business days after full execution of the Purchase Agreement')
add_bullet('Held in a mutually acceptable interest-bearing escrow account')
add_bullet('Refundable to Purchaser prior to Purchaser\'s election to proceed following initial due diligence and satisfaction of entitlement contingency conditions')
add_bullet('Non-refundable thereafter, except in the event of: (a) Seller default; (b) failure of the entitlement contingency; (c) failure of any condition to closing')

add_heading('4. Seller Financing')
p = doc.add_paragraph()
run = p.add_run('Seller shall carry a promissory note in the amount of $860,000 bearing interest at [4-5]% per annum, with the following terms:')
run.font.size = Pt(11)

add_bullet('Interest-only payments annually, commencing 12 months from closing')
add_bullet('Principal due at the earlier of: (a) simultaneous close with a builder/developer buyer; (b) sale of the Property or any portion thereof; (c) 24 months from the Closing Date')
add_bullet('No prepayment penalty')
add_bullet('Security: Purchase money mortgage or deed of trust on the Property')

add_heading('5. No Financing Contingency')
p = doc.add_paragraph()
run = p.add_run('Purchaser is prepared to pay all cash at closing using the seller carry structure. Purchaser\'s obligation is not conditioned on obtaining third-party financing.')
run.font.size = Pt(11)

add_heading('6. Closing Date and Timeline')
p = doc.add_paragraph()
run = p.add_run('Entitlement Period:').bold = True
run = p.add_run(' 24 months from Purchase Agreement execution — due diligence and entitlement process run concurrently throughout.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Closing Date:').bold = True
run = p.add_run(' Upon the earlier of: entitlement approval, builder simultaneous close, or 24 months from Purchase Agreement execution.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Extension Rights:').bold = True
run = p.add_run(' Purchaser may extend closing by two (2) additional 6-month periods upon payment of $25,000 per extension to Seller, credited to Purchase Price at closing.')
run.font.size = Pt(11)

add_heading('7. Due Diligence Investigation')
p = doc.add_paragraph()
run = p.add_run('Purchaser shall have access to the Property to investigate (at Purchaser\'s sole cost) throughout the Entitlement Period. Planned diligence items include: civil engineer concept plan / lot yield study, wetland delineation, topographic survey, geotechnical evaluation, utility feasibility, Phase I environmental assessment, title commitment / survey review, zoning and future land use verification, pre-application meeting with Greenville planning staff, review of Sub-Area Plan C documents, and review of Greenville Water & Wastewater Master Plans.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Purchaser may terminate the Agreement for any reason prior to electing to proceed with the entitlement process, with full return of the Earnest Money Deposit.')
run.font.size = Pt(11)

add_heading('8. Entitlement Contingency')
p = doc.add_paragraph()
run = p.add_run('Purchaser\'s obligation to close is expressly conditioned upon the Village of Greenville approving a subdivision plat, rezoning, or development plan for the Property that is commercially viable in Purchaser\'s reasonable discretion, including without limitation:')
run.font.size = Pt(11)

add_bullet('Approval of a subdivision layout with lot sizes and densities consistent with the Village\'s Sub-Area Plan C and applicable zoning')
add_bullet('Reasonable assurance that municipal sewer and water can serve the development at commercially feasible costs')
add_bullet('No material adverse change in zoning, future land use designation, or Sub-Area Plan applicable to the Property')

p = doc.add_paragraph()
run = p.add_run('The final lot yield shall be determined through the engineer\'s concept plan and the Village\'s approval process. The purchase price set forth in Section 2 shall be subject to adjustment based on the final approved lot count, to be negotiated in good faith between the parties at that time.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('If the entitlement contingency is not satisfied within 24 months (plus extensions), Purchaser may terminate the Purchase Agreement and receive return of the Earnest Money Deposit.')
run.font.size = Pt(11)

add_heading('9. Seller Cooperation')
p = doc.add_paragraph()
run = p.add_run('Seller agrees to: (a) sign all applications, plats, certified survey maps, and development agreements required for the entitlement process; (b) allow Purchaser and its consultants access to the Property for surveys, soils testing, wetland delineation, and other due diligence; (c) execute any necessary authorizations for utility connection and road access; (d) not unreasonably withhold, condition, or delay consent to applications or permits.')
run.font.size = Pt(11)

add_heading('10. Assignment')
p = doc.add_paragraph()
run = p.add_run('Purchaser shall have the right to assign this Agreement, in whole or in part, to: (a) an entity controlled by or under common control with Purchaser; (b) a builder, developer, or joint venture partner; (c) any third party, with Seller\'s consent not to be unreasonably withheld. No assignment shall relieve Purchaser of its obligations under this Agreement unless expressly agreed in writing by Seller.')
run.font.size = Pt(11)

add_heading('11. Closing Costs')
add_bullet('Title commitment and policy: Purchaser')
add_bullet('Survey: Purchaser')
add_bullet('Deed recording fees: Split equally')
add_bullet('Transfer tax: Split equally')
add_bullet('Legal fees: Each party pays its own')
add_bullet('Escrow/closing fee: Split equally')

add_heading('12. Representations and Warranties')
p = doc.add_paragraph()
run = p.add_run('Seller shall represent and warrant in the Purchase Agreement: (a) Seller has good and marketable title to the Property; (b) no pending or threatened legal proceedings affecting the Property; (c) no leases, tenancies, or third-party rights affecting the Property; (d) no environmental contamination known to Seller; (e) all real estate taxes are current; (f) Seller has authority to sell the Property and execute the Agreement.')
run.font.size = Pt(11)

add_heading('13. Access and Indemnification')
p = doc.add_paragraph()
run = p.add_run('Purchaser and its consultants may enter the Property for due diligence upon 24 hours\' notice to Seller. Purchaser shall indemnify Seller for any physical damage caused by such entry.')
run.font.size = Pt(11)

add_heading('14. Exclusive Negotiations')
p = doc.add_paragraph()
run = p.add_run('Seller agrees not to solicit, entertain, or negotiate with any third party regarding the sale, option, or development of the Property for the entire duration that this Agreement is in effect, including any extension periods. This exclusivity obligation shall terminate only upon the earlier of: (a) mutual written agreement of the parties; or (b) termination of this Agreement in accordance with its terms.')
run.font.size = Pt(11)

add_heading('15. Confidentiality')
p = doc.add_paragraph()
run = p.add_run('The terms of this LOI shall remain confidential between the parties and their respective advisors.')
run.font.size = Pt(11)

add_heading('16. Nonbinding')
p = doc.add_paragraph()
run = p.add_run('This LOI is a nonbinding expression of interest. No legal obligation shall exist until a definitive Purchase Agreement is executed by both parties. This paragraph and the Confidentiality and Exclusivity paragraphs are binding.')
run.font.size = Pt(11)

# Closing
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('If the foregoing terms are acceptable, please execute and return the counter-signed copy.')
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Very truly yours,')
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('__________________________')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Jason Christiansen')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Individually')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
run = p.add_run('Date: _______________')
run.font.size = Pt(11)

# Signature page break
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ACCEPTANCE')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('AGREED TO AND ACCEPTED this ______ day of ________________, 20____:')
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Estler/Barbara Family Trust (or Seller Entity)')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('__________________________')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Authorized Signatory')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
run = p.add_run('Date: _______________')
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('cc: Kevin [LAST NAME], [BROKERAGE]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
run = p.add_run('Purchaser\'s Counsel')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

# Save
output_path = '/root/wisconsin-overlay-map/output/Greenville_LOI.docx'
doc.save(output_path)
file_size = os.path.getsize(output_path)
print(f"LOI saved: {output_path} ({file_size:,} bytes)")